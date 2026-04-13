from pathlib import Path
from tempfile import TemporaryDirectory

from django.test import SimpleTestCase
from docx import Document as DocxDocument

from apps.ingestion.management.commands.generate_sample_inputs import build_simple_pdf
from apps.ingestion.mime_types import DOCX_MIME, PDF_MIME, detect_mime_type
from apps.ingestion.parsers.registry import parser_registry


class ParserTests(SimpleTestCase):
    def test_pdf_parser_extracts_text(self):
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "resume.pdf"
            file_path.write_bytes(
                build_simple_pdf(
                    [
                        "Alice Martin",
                        "Python Django Celery PostgreSQL",
                    ]
                )
            )

            mime_type = detect_mime_type(str(file_path))
            parser = parser_registry.get_parser(mime_type)
            parsed_document = parser.parse(str(file_path))

            self.assertEqual(mime_type, PDF_MIME)
            self.assertIn("Alice Martin", parsed_document.text)
            self.assertEqual(parsed_document.metadata["parser"], "pypdf")

    def test_docx_parser_extracts_text(self):
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "resume.docx"
            document = DocxDocument()
            document.add_paragraph("Ben Shah")
            document.add_paragraph("Embeddings and retrieval")
            document.save(file_path)

            mime_type = detect_mime_type(str(file_path))
            parser = parser_registry.get_parser(mime_type)
            parsed_document = parser.parse(str(file_path))

            self.assertEqual(mime_type, DOCX_MIME)
            self.assertIn("Ben Shah", parsed_document.text)
            self.assertEqual(parsed_document.metadata["parser"], "python-docx")
