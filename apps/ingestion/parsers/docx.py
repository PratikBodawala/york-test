from docx import Document as DocxDocument

from .base import BaseDocumentParser, ParsedDocument


class DocxDocumentParser(BaseDocumentParser):
    parser_name = "python-docx"

    def parse(self, file_path: str) -> ParsedDocument:
        document = DocxDocument(file_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text]
        text = "\n".join(paragraphs)
        return ParsedDocument(
            text=text,
            metadata={
                "paragraph_count": len(paragraphs),
                "parser": self.parser_name,
            },
        )
