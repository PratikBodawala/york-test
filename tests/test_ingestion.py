from pathlib import Path

from docx import Document as DocxDocument

from apps.ingestion.management.commands.generate_sample_inputs import build_simple_pdf
from apps.ingestion.mime_types import DOCX_MIME, PDF_MIME, detect_mime_type
from apps.ingestion.parsers.registry import parser_registry


def test_detect_and_parse_pdf(tmp_path: Path) -> None:
    pdf_path = tmp_path / "resume.pdf"
    pdf_path.write_bytes(
        build_simple_pdf(
            [
                "Alice Martin",
                "Senior Backend Engineer",
                "Python Django Celery Redis PostgreSQL",
            ]
        )
    )

    mime_type = detect_mime_type(str(pdf_path))
    parsed = parser_registry.get_parser(mime_type).parse(str(pdf_path))

    assert mime_type == PDF_MIME
    assert "Alice Martin" in parsed.text
    assert parsed.metadata["page_count"] == 1


def test_detect_and_parse_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "resume.docx"
    document = DocxDocument()
    document.add_paragraph("Ben Shah")
    document.add_paragraph("Data Platform Engineer")
    document.add_paragraph("Python Postgres embeddings batch processing")
    document.save(docx_path)

    mime_type = detect_mime_type(str(docx_path))
    parsed = parser_registry.get_parser(mime_type).parse(str(docx_path))

    assert mime_type == DOCX_MIME
    assert "Ben Shah" in parsed.text
    assert parsed.metadata["paragraph_count"] == 3
